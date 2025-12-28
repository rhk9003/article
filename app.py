import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

# --- 設定頁面配置 ---
st.set_page_config(layout="wide", page_title="文章修訂協作工具 v3", page_icon="📝")

# --- CSS 優化 ---
st.markdown("""
<style>
    .stCheckbox { padding-top: 10px; }
    .element-container { margin-bottom: -10px; }
</style>
""", unsafe_allow_html=True)

# --- Session State 初始化 ---
if 'doc_data' not in st.session_state:
    st.session_state['doc_data'] = [] 
if 'revisions' not in st.session_state:
    st.session_state['revisions'] = [] 
if 'next_id' not in st.session_state:
    st.session_state['next_id'] = 1
# 新增：儲存原始檔案的二進位資料，確保匯出時格式不流失
if 'original_file_bytes' not in st.session_state:
    st.session_state['original_file_bytes'] = None
if 'original_filename' not in st.session_state:
    st.session_state['original_filename'] = ""
# 新增：整體修改建議
if 'global_feedback' not in st.session_state:
    st.session_state['global_feedback'] = ""

# --- 輔助函式 ---

def read_file(uploaded_file):
    """
    讀取檔案內容用於預覽 (Preview)，同時回傳原始 bytes 用於最後匯出
    """
    # 讀取並儲存原始 bytes
    file_bytes = uploaded_file.getvalue()
    filename = uploaded_file.name
    
    doc_data = []
    
    if filename.endswith('.docx'):
        # 使用 BytesIO 讀取，不影響原始 bytes
        doc = Document(BytesIO(file_bytes))
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip(): 
                # 簡單判斷 Markdown 樣式用於預覽
                style_name = p.style.name
                md_prefix = ""
                if 'Heading 1' in style_name: md_prefix = "# "
                elif 'Heading 2' in style_name: md_prefix = "## "
                elif 'Heading 3' in style_name: md_prefix = "### "
                elif 'List Bullet' in style_name: md_prefix = "* "
                elif 'List Number' in style_name: md_prefix = "1. "
                
                doc_data.append({
                    'id': i,
                    'text': p.text,
                    'display_text': md_prefix + p.text,
                    'raw_text': p.text
                })
    elif filename.endswith('.txt'):
        stringio = file_bytes.decode("utf-8")
        lines = stringio.split('\n')
        for i, line in enumerate(lines):
            if line.strip():
                doc_data.append({
                    'id': i,
                    'text': line,
                    'display_text': line,
                    'raw_text': line
                })
        
    return doc_data, file_bytes, filename

def generate_appended_report(original_bytes, filename, global_feedback, revisions):
    """
    核心邏輯：讀取原始檔案 -> 在後方追加 (Append) 修改報告
    """
    if filename.endswith('.docx'):
        # 載入原始 Word 檔 (保留所有格式)
        doc = Document(BytesIO(original_bytes))
    else:
        # 如果是 txt，建立一個新的 Word 檔並填入內容
        doc = Document()
        doc.add_heading('原始文字內容', level=1)
        doc.add_paragraph(original_bytes.decode("utf-8"))

    # --- 開始追加內容 ---
    doc.add_page_break() # 強制換頁
    
    # 標題區
    doc.add_heading('【修訂建議報告】', level=0)
    
    # 1. 整體修改建議
    doc.add_heading('一、整體修改建議', level=1)
    if global_feedback.strip():
        doc.add_paragraph(global_feedback)
    else:
        doc.add_paragraph("（無整體建議）")
    
    # 2. 針對性修改 (表格呈現)
    doc.add_heading('二、細部修訂清單', level=1)
    
    if not revisions:
        doc.add_paragraph("無針對性修訂內容。")
    else:
        # 建立表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid' # Word 內建格線樣式
        
        # 設定表頭
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '編號'
        hdr_cells[1].text = '原始選取文字 (Target)'
        hdr_cells[2].text = '修改指示/建議 (Instruction)'
        
        # 填入內容
        for rev in revisions:
            row_cells = table.add_row().cells
            row_cells[0].text = str(rev['id'])
            row_cells[1].text = rev['target']
            row_cells[2].text = rev['instruction']
            
    # 儲存到 Buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 側邊欄 Sidebar ---

st.sidebar.title("🛠️ 修訂導航")

# 1. 整體修改建議區 (Global Feedback)
with st.sidebar.expander("📝 整體修改建議 (Global)", expanded=True):
    st.session_state['global_feedback'] = st.text_area(
        "請輸入對整篇文章的建議：",
        value=st.session_state['global_feedback'],
        height=150,
        placeholder="例如：文章語氣稍嫌生硬，建議多用主動語態..."
    )

st.sidebar.markdown("---")

# 2. 導航模式切換
mode_options = ["📄 閱讀與選取模式"]
for rev in st.session_state['revisions']:
    preview = (rev['target'][:15] + '..') if len(rev['target']) > 15 else rev['target']
    mode_options.append(f"#{rev['id']} 修訂: {preview}")

selection = st.sidebar.radio("功能選單：", mode_options)

st.sidebar.markdown("---")

# 3. 下載按鈕 (Append Logic)
if st.session_state['original_file_bytes']:
    st.sidebar.header("📤 匯出")
    st.sidebar.caption("將會輸出原始檔案（保留格式）並在文末附上修訂表。")
    
    docx_file = generate_appended_report(
        st.session_state['original_file_bytes'], 
        st.session_state['original_filename'],
        st.session_state['global_feedback'],
        st.session_state['revisions']
    )
    
    output_name = f"Revised_{st.session_state['original_filename']}" if st.session_state['original_filename'].endswith('.docx') else "Revised_Report.docx"
    
    st.sidebar.download_button(
        label="📥 下載完整 Word 報告",
        data=docx_file,
        file_name=output_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- 主頁面 Main Area ---

if not st.session_state['doc_data']:
    st.header("1. 上傳文章")
    st.info("支援 .docx (匯出時將完美保留原始格式) 與 .txt")
    uploaded_file = st.file_uploader("請上傳檔案", type=['docx', 'txt'])
    
    if uploaded_file is not None:
        # 讀取並存入 Session State
        data, file_bytes, filename = read_file(uploaded_file)
        st.session_state['doc_data'] = data
        st.session_state['original_file_bytes'] = file_bytes
        st.session_state['original_filename'] = filename
        st.rerun()

else:
    # --- 閱讀與選取模式 ---
    if selection == "📄 閱讀與選取模式":
        st.title("文章閱讀與標記")
        
        # 顯示目前的整體建議 (唯讀預覽)
        if st.session_state['global_feedback']:
            st.info(f"💡 目前的整體建議：{st.session_state['global_feedback']}")
        
        col_main, col_action = st.columns([3, 1])
        
        with col_action:
            # 浮動操作區
            st.markdown("### ⚡ 操作區")
            st.caption("勾選左側段落後加入修訂：")
            
            if st.button("➕ 將勾選段落加入修訂", type="primary"):
                selected_texts = []
                for item in st.session_state['doc_data']:
                    key = f"chk_{item['id']}"
                    if st.session_state.get(key, False):
                        selected_texts.append(item['raw_text'])
                        st.session_state[key] = False # 重置勾選
                
                if selected_texts:
                    combined_text = "\n\n".join(selected_texts)
                    new_rev = {
                        'id': st.session_state['next_id'],
                        'target': combined_text,
                        'instruction': "" 
                    }
                    st.session_state['revisions'].append(new_rev)
                    st.session_state['next_id'] += 1
                    st.success(f"已建立修訂 #{new_rev['id']}")
                    st.rerun()
                else:
                    st.warning("請先勾選段落！")

            if st.button("🧹 清除所有勾選"):
                 for item in st.session_state['doc_data']:
                    key = f"chk_{item['id']}"
                    if key in st.session_state:
                        st.session_state[key] = False
                 st.rerun()
            
            st.markdown("---")
            if st.button("🔄 上傳新文件 (重置)"):
                st.session_state.clear()
                st.rerun()

        with col_main:
            st.subheader("文件內容")
            for item in st.session_state['doc_data']:
                c1, c2 = st.columns([0.5, 9.5])
                with c1:
                    st.checkbox("", key=f"chk_{item['id']}")
                with c2:
                    st.markdown(item['display_text'])

    # --- 編輯修訂模式 ---
    else:
        selected_id = int(selection.split(" ")[0].replace("#", ""))
        current_rev = next((item for item in st.session_state['revisions'] if item['id'] == selected_id), None)
        
        if current_rev:
            st.title(f"編輯修訂項目 #{selected_id}")
            
            st.caption("原始選取文字：")
            st.text_area("Target", value=current_rev['target'], height=150, disabled=True)
            
            st.subheader("👇 修改建議")
            new_instruction = st.text_area(
                "請輸入具體修改指示...", 
                value=current_rev['instruction'], 
                height=200,
                key=f"inst_{selected_id}"
            )
            
            col_save, col_del = st.columns([1, 4])
            with col_save:
                if st.button("💾 儲存內容"):
                    for item in st.session_state['revisions']:
                        if item['id'] == selected_id:
                            item['instruction'] = new_instruction
                    st.success("已儲存！")
            
            with col_del:
                if st.button("🗑️ 刪除此修訂"):
                    st.session_state['revisions'] = [item for item in st.session_state['revisions'] if item['id'] != selected_id]
                    st.rerun()
